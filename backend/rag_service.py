"""RAG service — extract text from PDF/DOCX, chunk, hybrid BM25 + semantic retrieve.

Retrieval dulu murni BM25 (lexical/keyword) - upgrade 2026-07-19 ke hybrid: kalau
OpenAI embeddings tersedia (pakai EMERGENT_LLM_KEY yang di deployment ini adalah key
OpenAI asli, lihat ai_service.py), skor semantik digabung dengan BM25 supaya pertanyaan
yang beda kata tapi sama makna ("berapa biaya kalau batal" vs "kebijakan refund") tetap
ketemu dokumennya. SENGAJA graceful degrade ke BM25 murni kalau embedding API gagal/tidak
dikonfigurasi - RAG tidak boleh berhenti total gara-gara satu provider down.
Chunks disimpan di MongoDB `rag_chunks` (field `embedding` opsional per chunk, None kalau
belum/gagal di-embed - chunk lama dari sebelum fitur ini tetap jalan lewat BM25 saja).
"""
import io
import logging
import os
import re
from typing import List, Optional

import httpx
import numpy as np
from pypdf import PdfReader
from docx import Document as DocxDocument
from rank_bm25 import BM25Okapi

EMBED_MODEL = "text-embedding-3-small"
EMBED_API_URL = "https://api.openai.com/v1/embeddings"
_embed_logger = logging.getLogger("rag_embeddings")


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(parts)


def extract_text_from_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_txt(data: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if name.endswith(".docx"):
        return extract_text_from_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return extract_text_from_txt(data)
    raise ValueError(f"Unsupported file type: {filename}")


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 80) -> List[str]:
    """Simple sliding window over cleaned paragraphs."""
    # normalise whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # try to end on sentence boundary
        if end < len(text):
            dot = text.rfind(". ", start, end)
            nl = text.rfind("\n", start, end)
            pivot = max(dot, nl)
            if pivot > start + chunk_size * 0.4:
                end = pivot + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


_TOKEN_RE = re.compile(r"[A-Za-zÀ-ÿ0-9]+", re.UNICODE)


def tokenize(s: str) -> List[str]:
    return [t.lower() for t in _TOKEN_RE.findall(s or "")]


def bm25_search(query: str, chunks: List[dict], k: int = 5) -> List[dict]:
    """Return top-k chunks with score. Each chunk is {id, doc_id, doc_title, text}."""
    if not chunks:
        return []
    corpus = [tokenize(c["text"]) for c in chunks]
    bm25 = BM25Okapi(corpus)
    q_tokens = tokenize(query)
    if not q_tokens:
        return []
    scores = bm25.get_scores(q_tokens)
    ranked = sorted(zip(scores, chunks), key=lambda x: x[0], reverse=True)
    out = []
    # keep top-k; only drop chunks with 0 keyword overlap
    q_set = set(q_tokens)
    for score, chunk in ranked[:k]:
        overlap = q_set.intersection(tokenize(chunk["text"]))
        if not overlap:
            continue
        out.append({**chunk, "score": float(max(score, 0.01))})
    return out


def _embedding_api_key() -> Optional[str]:
    key = os.environ.get("EMERGENT_LLM_KEY", "")
    return key if key.startswith("sk-") and not key.startswith("sk-emergent-") else None


async def get_embeddings_batch(texts: List[str]) -> Optional[List[Optional[List[float]]]]:
    """Batch embed lewat OpenAI embeddings API langsung (bukan lewat emergentintegrations -
    SDK itu tidak punya kelas embeddings). Return None (bukan list) kalau API key belum
    ada/dikonfigurasi sama sekali - beda dari "berhasil tapi kosong" supaya caller tahu
    harus fallback ke BM25, bukan menganggap semua chunk sengaja tidak relevan."""
    api_key = _embedding_api_key()
    if not api_key or not texts:
        return None
    try:
        async with httpx.AsyncClient(timeout=30) as http:
            resp = await http.post(
                EMBED_API_URL,
                headers={"Authorization": f"Bearer {api_key}"},
                json={"model": EMBED_MODEL, "input": texts},
            )
        if resp.status_code >= 400:
            _embed_logger.warning(f"Embedding API gagal HTTP {resp.status_code}: {resp.text[:300]}")
            return None
        data = resp.json()
        by_index = {item["index"]: item["embedding"] for item in data["data"]}
        return [by_index.get(i) for i in range(len(texts))]
    except Exception as e:
        _embed_logger.warning(f"Gagal memanggil embedding API: {e}")
        return None


async def get_embedding(text: str) -> Optional[List[float]]:
    result = await get_embeddings_batch([text])
    return result[0] if result else None


def _cosine(a: List[float], b: List[float]) -> float:
    va, vb = np.array(a), np.array(b)
    denom = (np.linalg.norm(va) * np.linalg.norm(vb)) or 1.0
    return float(np.dot(va, vb) / denom)


async def hybrid_search(query: str, chunks: List[dict], k: int = 5) -> List[dict]:
    """BM25 + semantic (kalau embedding tersedia), digabung lewat weighted sum skor yang
    sudah dinormalisasi 0-1 masing-masing (0.4 BM25 + 0.6 semantic untuk chunk yang punya
    embedding; chunk tanpa embedding cuma dapat skor BM25 - tidak pernah didiskualifikasi,
    cuma tidak dapat boost semantik). Kalau query gagal di-embed sama sekali (API down/key
    belum ada), otomatis 100% BM25 - PERSIS perilaku lama, tidak ada regresi."""
    if not chunks:
        return []
    q_tokens = tokenize(query)
    if not q_tokens:
        return []

    corpus = [tokenize(c["text"]) for c in chunks]
    bm25 = BM25Okapi(corpus)
    bm25_scores = bm25.get_scores(q_tokens)
    bm25_max = max(bm25_scores) or 1.0
    bm25_norm = [s / bm25_max for s in bm25_scores]

    q_embedding = await get_embedding(query)
    semantic_norm = [0.0] * len(chunks)
    if q_embedding:
        sims = [_cosine(q_embedding, c["embedding"]) if c.get("embedding") else 0.0 for c in chunks]
        sem_max = max(sims) or 1.0
        semantic_norm = [max(s, 0.0) / sem_max for s in sims]

    q_set = set(q_tokens)
    scored = []
    for i, chunk in enumerate(chunks):
        has_embedding = q_embedding and chunk.get("embedding")
        combined = (0.4 * bm25_norm[i] + 0.6 * semantic_norm[i]) if has_embedding else bm25_norm[i]
        overlap = q_set.intersection(tokenize(chunk["text"]))
        # chunk relevan kalau ada overlap kata KUNCI ATAU skor semantik tinggi (>0.5) -
        # ini yang bikin hybrid lebih pintar dari BM25 murni: query berbeda kata tapi
        # semantiknya dekat tetap lolos walau tidak ada overlap literal.
        if not overlap and semantic_norm[i] < 0.5:
            continue
        scored.append((combined, chunk))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [{**c, "score": float(max(s, 0.01))} for s, c in scored[:k]]


def build_rag_context(hits: List[dict], max_chars: int = 3500) -> str:
    if not hits:
        return ""
    out = ["=== DOKUMEN REFERENSI (RAG) ==="]
    total = 0
    for h in hits:
        block = f"[{h.get('doc_title','doc')}] {h['text']}"
        if total + len(block) > max_chars:
            break
        out.append(block)
        total += len(block)
    out.append("=== AKHIR DOKUMEN ===")
    return "\n\n".join(out)
